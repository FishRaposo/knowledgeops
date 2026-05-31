"""DOCX document parser using python-docx."""

from typing import BinaryIO

from docx import Document

from app.parsers.base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Parser for DOCX documents using python-docx.

    Extracts text from paragraphs, handling headers and tables.
    """

    def parse(self, file: BinaryIO, filename: str) -> ParseResult:
        """Parse a DOCX file and extract text content.

        Args:
            file: DOCX file binary stream.
            filename: Original filename.

        Returns:
            ParseResult with text from all paragraphs and tables.
        """
        doc = Document(file)

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        title = paragraphs[0] if paragraphs else filename.rsplit(".", 1)[0]
        content = "\n\n".join(paragraphs)

        return ParseResult(
            title=title,
            content=content,
            metadata={
                "format": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "char_count": len(content),
            },
        )
